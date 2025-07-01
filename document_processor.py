import os
import tempfile
from typing import Union, Any
import streamlit as st

# Document processing libraries
try:
    import PyPDF2
except ImportError:
    try:
        import pypdf as PyPDF2
    except ImportError:
        PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

class DocumentProcessor:
    """Handles processing of various document formats"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def process_file(self, uploaded_file: Any) -> str:
        """Process uploaded file and extract text content"""
        
        if uploaded_file is None:
            raise ValueError("No file provided")
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            tmp_file_path = tmp_file.name
        
        try:
            if file_extension == 'pdf':
                content = self._process_pdf(tmp_file_path)
            elif file_extension == 'docx':
                content = self._process_docx(tmp_file_path)
            elif file_extension == 'txt':
                content = self._process_txt(tmp_file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            return content
            
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        
        if PyPDF2 is None:
            raise ImportError("PyPDF2 or pypdf library is required for PDF processing")
        
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                        continue
                
                if not text_content:
                    raise ValueError("No text content could be extracted from the PDF")
                
                return "\n\n".join(text_content)
                
        except Exception as e:
            raise ValueError(f"Error processing PDF file: {str(e)}")
    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        
        if Document is None:
            raise ImportError("python-docx library is required for DOCX processing")
        
        try:
            doc = Document(file_path)
            
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("--- Table ---\n" + "\n".join(table_text))
            
            if not text_content:
                raise ValueError("No text content could be extracted from the DOCX file")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise ValueError(f"Error processing DOCX file: {str(e)}")
    
    def _process_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                        if content.strip():
                            return content
                        else:
                            raise ValueError("Text file is empty")
                            
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode the text file with any supported encoding")
            
        except Exception as e:
            raise ValueError(f"Error processing text file: {str(e)}")
    
    def get_file_info(self, uploaded_file: Any) -> dict:
        """Get information about uploaded file"""
        
        if uploaded_file is None:
            return {}
        
        return {
            'name': uploaded_file.name,
            'size': uploaded_file.size,
            'type': uploaded_file.type,
            'extension': uploaded_file.name.split('.')[-1].lower() if '.' in uploaded_file.name else 'unknown'
        }
    
    def validate_file(self, uploaded_file: Any) -> tuple[bool, str]:
        """Validate uploaded file"""
        
        if uploaded_file is None:
            return False, "No file provided"
        
        file_info = self.get_file_info(uploaded_file)
        
        # Check file extension
        if file_info['extension'] not in self.supported_formats:
            return False, f"Unsupported file format: {file_info['extension']}"
        
        # Check file size (limit to 10MB)
        max_size = 10 * 1024 * 1024  # 10MB
        if file_info['size'] > max_size:
            return False, f"File size ({file_info['size']} bytes) exceeds maximum allowed size ({max_size} bytes)"
        
        return True, "File is valid"
